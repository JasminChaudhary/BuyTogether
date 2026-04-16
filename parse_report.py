import sys
try:
    from docx import Document
except ImportError:
    print("python-docx not installed.")
    sys.exit(1)

doc = Document('Final_Report.docx')

with open('extracted_report.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            f.write(f"P{i}: {text}\n")
    
    f.write("\n\n--- TABLES ---\n\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"Table {t_idx}:\n")
        for row in table.rows:
            f.write(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]) + "\n")
        f.write("\n")

print("Report extracted successfully.")
